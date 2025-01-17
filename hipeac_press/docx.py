import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from .transformers.html import HtmlTransformer
from .transformers.markdown import MarkdownTransformer
from .transformers.pdf import PdfTransformer
from .type_definitions import (
    Author,
    BulletList,
    Document,
    Header,
    Image,
    NavItem,
    OrderedList,
    Paragraph,
    Quote,
    Reference,
)
from .utils.slug import slugify


class DocxConverter:
    """A class to convert DOCX files to a structured document format."""

    def __init__(
        self,
        docx_path: Path,
        *,
        img_folder: Path,
        metadata_path: Path = None,
        section_name: str = None,
        prev: NavItem | None = None,
        next: NavItem | None = None,
    ):
        """Initialize the DocxConverter with paths and metadata.

        :param docx_path: Path to the DOCX file.
        :param img_folder: Directory to save images.
        :param metadata_path: Path to the metadata JSON file.
        """
        self._docx_path = docx_path
        self._docx = DocxDocument(docx_path)
        self.errors = []  # Add error tracking
        self.section_name = section_name

        self.img_folder = self._generate_img_folder(img_folder)
        self.metadata = self._read_metadata(metadata_path)
        self.document = self._create_document(prev=prev, next=next)

    def _generate_img_folder(self, base_folder: Path) -> Path:
        """Generate a unique image directory based on the DOCX path.

        :param base_folder: Base directory for images.
        :returns: The generated image directory.
        """
        base_folder.mkdir(parents=True, exist_ok=True)  # Ensure the base directory exists
        hash_object = hashlib.md5(str(self._docx_path).encode())
        img_folder = base_folder / hash_object.hexdigest()
        img_folder.mkdir(parents=True, exist_ok=True)

        return img_folder

    def _read_metadata(self, metadata_path: Path) -> dict:
        """Read metadata from a JSON file or the DOCX core properties.

        :param metadata_path: Path to the metadata JSON file.
        :returns: The metadata dictionary.
        """
        core_properties = self._docx.core_properties
        metadata = {
            "title": core_properties.title,
            "authors": core_properties.author.split(",") if core_properties.author else [],
            "keywords": core_properties.keywords.split(",") if core_properties.keywords else [],
        }

        if metadata_path and metadata_path.exists():
            with open(metadata_path) as file:
                json_metadata = json.load(file)
                metadata.update(json_metadata)

        return metadata

    def _save_image(self, run) -> Path | None:
        """Save an image from the DOCX run to the image directory.

        :param run: The DOCX run containing the image.
        :returns: The path to the saved image or None if the image cannot be saved.
        """
        blip = run._element.xpath(".//a:blip")[0]
        rId = blip.get(qn("r:embed"))
        image_part = self._docx.part.related_parts[rId]
        image_filename = Path(image_part.partname).name

        # Skip EMF and WMF files
        for ext in [".emf", ".wmf"]:
            if image_filename.lower().endswith(ext):
                self.errors.append(f"{ext} image format not supported: {image_filename}")
                return None

        img_path = self.img_folder / image_filename

        try:
            with open(img_path.absolute(), "wb") as img_file:
                img_file.write(image_part.blob)
        except Exception as e:
            self.errors.append(f"Failed to save image {image_filename}: {e}")
            return None

        return img_path

    def _convert_run_text(self, run) -> str:
        """Convert a run's text with its formatting.

        :param run: The DOCX run to convert.
        :returns: The formatted text.
        """
        text = run.text
        if text.strip():
            if run.bold and run.italic:
                text = f"**_{text}_**"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"_{text}_"
        return text

    def _consolidate_formatting(self, text: str) -> str:
        """Consolidate adjacent formatting markers and trim results."""
        # First, consolidate simple adjacent markers
        text = re.sub(r"\*\*([^*]+)\*\*\*\*([^*]+)\*\*", r"**\1\2**", text)
        text = re.sub(r"_([^_]+)__([^_]+)_", r"_\1\2_", text)

        # Then handle more complex patterns
        # Find all bold sections and consolidate them
        while True:
            # Look for adjacent bold sections like "**abc****def**"
            pattern = r"\*\*[^*]+\*\*\*\*[^*]+\*\*"
            match = re.search(pattern, text)
            if not match:
                break

            # Extract the content between ** markers
            section = match.group(0)
            content = re.findall(r"\*\*([^*]+)\*\*", section)
            # Combine the content and wrap in single **
            combined = f"**{''.join(content)}**"
            # Replace the original section with the combined version
            text = text.replace(section, combined)

        # Trim spaces inside formatting markers
        text = re.sub(r"\*\*\s*([^*]+?)\s*\*\*", r"**\1**", text)
        text = re.sub(r"_\s*([^_]+?)\s*_", r"_\1_", text)

        return text

    def _get_list_type(self, paragraph) -> str | None:
        """Determine if a paragraph is a list item and its type by checking the numFmt attribute."""
        # Get numbering properties
        pPr = paragraph._element.xpath("./w:pPr/w:numPr")

        if not pPr:
            return None

        # Get the numbering id
        num_id = pPr[0].find(qn("w:numId")).get(qn("w:val"))

        # Get the formatting from the numbering definitions
        numbering = self._docx.part.numbering_part.numbering_definitions._numbering
        num = numbering.xpath(f"//w:num[@w:numId='{num_id}']")
        if not num:
            return "bullet"

        ilvl = pPr[0].find(qn("w:ilvl"))
        ilvl = ilvl.get(qn("w:val")) if ilvl is not None else "0"

        abstract_num_id = num[0].find(qn("w:abstractNumId")).get(qn("w:val"))
        num_fmt = numbering.xpath(
            f"//w:abstractNum[@w:abstractNumId='{abstract_num_id}']/w:lvl[@w:ilvl='{ilvl}']/w:numFmt"
        )[0].get(qn("w:val"))

        # According to OOXML spec, these formats indicate ordered lists
        ordered_formats = {"decimal", "upperRoman", "lowerRoman", "upperLetter", "lowerLetter"}

        return "ordered" if num_fmt in ordered_formats else "bullet"

    def _convert_paragraphs(self):
        """Convert paragraphs in the DOCX to structured document elements."""
        elements = []
        references = []
        in_references_section = False
        current_list_items = []
        current_list_type = None
        last_processed_para_index = -1  # Initialize the index

        for i, para in enumerate(self._docx.paragraphs):
            if i <= last_processed_para_index:
                continue

            # Check if we're in the references section
            if para.text.strip().lower() == "references":
                in_references_section = True
                continue

            # If we're in references section, collect references
            if in_references_section:
                if para.text.strip():
                    references.append(Paragraph(text=para.text.strip()))
                continue

            # Check for lists using both style name and XML structure
            is_list = para._p.pPr is not None and para._p.pPr.numPr is not None

            if is_list:
                # Get the list type for this paragraph
                list_type = self._get_list_type(para)

                # If this is the first item or list type changes, flush current list
                if current_list_items and list_type != current_list_type:
                    if current_list_type == "bullet":
                        elements.append(BulletList(items=current_list_items))
                    else:
                        elements.append(OrderedList(items=current_list_items))
                    current_list_items = []

                # Set or update list type
                current_list_type = list_type or "bullet"

                # Process list item text
                formatted_text = ""
                for run in para.runs:
                    formatted_text += self._convert_run_text(run)
                formatted_text = self._consolidate_formatting(formatted_text.strip())
                current_list_items.append(formatted_text)
                continue

            # Not a list item - flush any pending list
            if current_list_items:
                if current_list_type == "bullet":
                    elements.append(BulletList(items=current_list_items))
                else:
                    elements.append(OrderedList(items=current_list_items))
                current_list_items = []
                current_list_type = None

            # Check for headers
            if para.style.name.startswith("Heading"):
                level = int(para.style.name.replace("Heading ", "")) or 1
                elements.append(Header(level=level, text=para.text.strip()))
                continue

            # Check for quotes
            if para.style.name == "Quote":
                quote_text = para.text.strip()
                ref = None
                if i + 1 < len(self._docx.paragraphs):
                    next_para = self._docx.paragraphs[i + 1]
                    ref = Paragraph(text=next_para.text.strip())
                    last_processed_para_index = i + 1  # Update the index
                elements.append(Quote(text=quote_text, ref=ref))
                continue

            # Normal paragraph processing
            paragraph_elements = []
            formatted_text = ""

            for run in para.runs:
                if run._element.xpath(".//a:blip"):
                    # Handle image
                    img_path = self._save_image(run)
                    if img_path:
                        if formatted_text.strip():
                            # Consolidate formatting before creating paragraph
                            formatted_text = self._consolidate_formatting(formatted_text.strip())
                            paragraph_elements.append(Paragraph(text=formatted_text))
                            formatted_text = ""
                        caption = None
                        if i + 1 < len(self._docx.paragraphs) and self._docx.paragraphs[i + 1].style.name == "Caption":
                            caption = self._docx.paragraphs[i + 1].text
                        paragraph_elements.append(Image(path=img_path, caption=caption))
                else:
                    # Handle formatted text
                    formatted_text += self._convert_run_text(run)

            if formatted_text.strip() and para.style.name != "Caption":
                # Consolidate formatting before creating paragraph
                formatted_text = self._consolidate_formatting(formatted_text.strip())
                paragraph_elements.append(Paragraph(text=formatted_text))

            elements.extend(paragraph_elements)

        # Flush any remaining list
        if current_list_items:
            if current_list_type == "bullet":
                elements.append(BulletList(items=current_list_items))
            else:
                elements.append(OrderedList(items=current_list_items))

        # clean up references

        refs = []

        for ref in references:
            try:
                code = ref.text.split("]")[0][1:]
                text = ref.text.split("]")[1].strip()
                refs.append(Reference(code=code, text=text))
            except Exception as e:
                self.errors.append(f"Failed to parse reference: {ref.text}, error: {e}")

        refs.sort(key=lambda x: x.code.lower())

        return elements, refs

    def _create_document(self, *, prev: NavItem | None = None, next: NavItem | None = None) -> Document:
        """Create a structured document from the DOCX content."""
        elements, references = self._convert_paragraphs()
        doc = Document(
            slug=f"{slugify(self.section_name or 'n')}--{slugify(self.metadata.get('title', ''))}",
            title=self.metadata.get("title", ""),
            authors=[Author(name=author) for author in self.metadata.get("authors", [])],
            keywords=self.metadata.get("keywords", []),
            updated_at=datetime.now(),
            prev=prev,
            next=next,
            references=references,  # Add references to the document
        )
        doc.elements.extend(elements)
        return doc

    def export(self, format: str = "md", **kwargs) -> bytes:
        """Export the structured document using the specified transformer.

        :param transformer: The transformer to use for exporting.
        :returns: The transformed document as a string.
        """
        if format == "html":
            return HtmlTransformer(self.document).get(v=kwargs.get("v", 5))
        elif format == "md":
            return MarkdownTransformer(self.document).get()
        elif format == "pdf":
            return PdfTransformer(self.document, kwargs.get("build_path")).get(section=kwargs.get("section"))
        else:
            raise ValueError(f"Unsupported format: {format}")

    @property
    def title(self) -> str:
        """Return the title of the document, or the filename if no title is found."""
        return self.document.title or self._docx_path.stem

    @property
    def slug(self) -> str:
        """Return a slug based on the section name and the title of the document."""
        return f"{slugify(self.section_name or 'n')}--{slugify(self.title)}"

    def add_element(self, element):
        """Add an element to the document."""
        self.document.elements.append(element)

    def add_reference(self, reference: Reference):
        """Add a reference to the document."""
        self.document.references.append(reference)

    def set_prev(self, prev: NavItem):
        """Set the previous document in the tree."""
        self.document.prev = prev

    def set_next(self, next: NavItem):
        """Set the next document in the tree."""
        self.document.next = next
